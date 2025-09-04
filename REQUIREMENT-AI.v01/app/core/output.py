from pathlib import Path
from typing import List, Dict

import pandas as pd
from docx import Document


def write_docx(reqs: List[Dict], path: str | Path) -> None:
    doc = Document()
    doc.add_heading('Software Requirements Specification', level=1)
    for r in reqs:
        p = doc.add_paragraph()
        p.add_run(f"[{r.get('moscow','')}] ").bold = True
        p.add_run(r["text"]) 
        p.add_run(f" ({r.get('category','')})")
    doc.save(str(path))


def write_excel(reqs: List[Dict], path: str | Path) -> None:
    df = pd.DataFrame(reqs)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(path, index=False)


def generate_user_stories(reqs: List[Dict]) -> List[str]:
    stories = []
    for r in reqs:
        stories.append(f"As a user, I want {r['text']} so that value is delivered.")
    return stories

def write_user_stories(stories: List[str], path: str | Path) -> None:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("\n".join(stories), encoding="utf-8")


